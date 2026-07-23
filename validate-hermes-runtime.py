#!/opt/hermes/.venv/bin/python
"""Secret-safe smoke test for the persistent Hermes runtime."""

from __future__ import annotations

import argparse
import importlib
import json
import os
from pathlib import Path
import shutil
import subprocess
import tempfile
import urllib.request

import yaml


parser = argparse.ArgumentParser()
parser.add_argument(
    "--network",
    action="store_true",
    help="also run a live key-free web search",
)
args = parser.parse_args()


COMMANDS = (
    "file",
    "hermes",
    "hermes-admin",
    "hermes-smoke-test",
    "jq",
    "libreoffice",
    "node",
    "npm",
    "pdfinfo",
    "pdftotext",
    "pip",
    "tesseract",
    "uv",
    "validate-vision",
)
MODULES = (
    "bs4",
    "docx",
    "ddgs",
    "fitz",
    "fastapi",
    "httpx",
    "lxml",
    "markitdown",
    "openpyxl",
    "pdfplumber",
    "pypdf",
    "uvicorn",
)


def run(*command: str, timeout: int = 20) -> str:
    result = subprocess.run(
        command,
        check=True,
        capture_output=True,
        text=True,
        timeout=timeout,
    )
    return result.stdout


checks: dict[str, bool] = {}
checks["non_root_user"] = os.getuid() == 10000
for command in COMMANDS:
    checks[f"command_{command}"] = shutil.which(command) is not None
for module in MODULES:
    importlib.import_module(module)
    checks[f"module_{module}"] = True

config = yaml.safe_load(Path("/opt/data/config.yaml").read_text(encoding="utf-8")) or {}
checks["smart_approvals"] = config.get("approvals", {}).get("mode") == "smart"
checks["cron_approvals"] = config.get("approvals", {}).get("cron_mode") == "approve"
checks["loop_hard_stop"] = (
    config.get("tool_loop_guardrails", {}).get("hard_stop_enabled") is True
)
checks["shell_profile"] = config.get("terminal", {}).get("shell_init_files") == [
    "/opt/data/home/.hermes_env"
]
checks["writable_terminal_cwd"] = (
    config.get("terminal", {}).get("cwd") == "/opt/data/home"
)
checks["serialized_cron"] = config.get("cron", {}).get("max_parallel_jobs") == 1
checks["fallbacks"] = len(config.get("fallback_providers") or []) >= 1
fallbacks = config.get("fallback_providers") or []
checks["local_gpu_first_fallback"] = bool(
    fallbacks
    and fallbacks[0].get("provider") == "ollama-local"
    and fallbacks[0].get("model") == "qwen3-4b-gpu:latest"
)
custom_providers = {
    item.get("name"): item
    for item in (config.get("custom_providers") or [])
    if isinstance(item, dict) and item.get("name")
}
checks["ollama_bridge_configured"] = (
    custom_providers.get("ollama-local", {}).get("base_url")
    == "http://ollama-bridge:8000/v1"
)
if os.getenv("NARAROUTER_API_KEY"):
    checks["nararouter_primary"] = (
        config.get("model", {}).get("provider") == "nararouter"
        and config.get("model", {}).get("default") == "mistral-large"
    )
    checks["nararouter_key_env"] = (
        custom_providers.get("nararouter", {}).get("key_env")
        == "NARAROUTER_API_KEY"
    )
checks["auxiliary_auto_routing"] = all(
    task == "vision"
    or not isinstance(settings, dict)
    or settings.get("provider") == "auto"
    for task, settings in (config.get("auxiliary") or {}).items()
)
delegation = config.get("delegation") or {}
expected_delegation_provider = (
    "nararouter" if os.getenv("NARAROUTER_API_KEY") else "ollama-local"
)
expected_delegation_model = (
    "mistral-large"
    if expected_delegation_provider == "nararouter"
    else "qwen3-4b-gpu:latest"
)
checks["delegation_route"] = (
    delegation.get("provider") == expected_delegation_provider
    and delegation.get("model") == expected_delegation_model
)
checks["delegation_completion_budget"] = (
    delegation.get("max_iterations") == 20
    and delegation.get("child_timeout_seconds") == 600
)
checks["delegation_concurrency_guard"] = (
    delegation.get("max_concurrent_children") == 1
    and delegation.get("max_spawn_depth") == 1
)
checks["delegation_summary_guard"] = delegation.get("max_summary_chars") == 12000
checks["ddgs_search"] = config.get("web", {}).get("search_backend") == "ddgs"
vision = config.get("auxiliary", {}).get("vision", {})
checks["vision_route"] = (
    (
        vision.get("provider") == "nararouter"
        and vision.get("model") == "mistral-medium-3-5"
    )
    or (
        vision.get("provider") == "openrouter"
        and vision.get("model") == "google/gemma-4-26b-a4b-it:free"
    )
    or (
        vision.get("provider") == "gemini"
        and vision.get("model") == "gemini-2.5-flash"
    )
)
vision_fallback = vision.get("fallback_chain") or []
if vision_fallback:
    checks["vision_fallback"] = any(
        entry.get("provider") == "gemini"
        and entry.get("model") == "gemini-2.5-flash"
        for entry in vision_fallback
        if isinstance(entry, dict)
    )
checks["provider_timeout"] = (
    config.get("providers", {}).get("openrouter", {}).get(
        "request_timeout_seconds"
    )
    == 60
)
checks["nararouter_stale_timeout"] = (
    config.get("providers", {}).get("nararouter", {}).get(
        "stale_timeout_seconds"
    )
    == 15
)
checks["gateway_timeout"] = config.get("agent", {}).get("gateway_timeout") == 300
checks["runtime_policy"] = "[HERMES_RUNTIME_POLICY]" in str(
    config.get("agent", {}).get("system_prompt") or ""
)
checks["general_agent_identity"] = "general-purpose personal assistant" in str(
    config.get("agent", {}).get("system_prompt") or ""
)
checks["delegation_policy"] = "Delegated workers use" in str(
    config.get("agent", {}).get("system_prompt") or ""
)
checks["quota_policy"] = "Use hermes-admin quota" in str(
    config.get("agent", {}).get("system_prompt") or ""
)
checks["artifact_delivery_policy"] = (
    "A path or MEDIA marker in prose is not proof of delivery"
    in str(config.get("agent", {}).get("system_prompt") or "")
)
if os.getenv("TELEGRAM_BOT_TOKEN"):
    checks["telegram_home_channel"] = bool(
        str(config.get("TELEGRAM_HOME_CHANNEL") or "").strip()
    )
checks["verify_on_stop"] = config.get("agent", {}).get("verify_on_stop") == "auto"
checks["skill_bundles"] = (
    len(list(Path("/opt/data/skill-bundles").glob("*.yaml"))) >= 5
)
try:
    version = run("hermes", "--version", timeout=10)
    checks["hermes_release"] = "v0.19.0" in version
except (OSError, subprocess.SubprocessError):
    checks["hermes_release"] = False

for raw_path in (
    "/opt/data/cache/uv",
    "/opt/data/cache/huggingface",
    "/opt/data/cache/pytest",
    "/opt/data/python-packages",
    "/opt/data/tmp",
):
    path = Path(raw_path)
    checks[f"owner_{path.name}"] = path.stat().st_uid == 10000
    checks[f"writable_{path.name}"] = os.access(path, os.W_OK)

with tempfile.TemporaryDirectory(prefix="hermes-runtime-") as raw_temp:
    temp = Path(raw_temp)

    import fitz

    pdf_path = temp / "sample.pdf"
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "HERMES PDF EXTRACTION OK")
    document.save(pdf_path)
    document.close()
    pdf_text = run("pdftotext", str(pdf_path), "-")
    checks["pdf_extract"] = "HERMES PDF EXTRACTION OK" in pdf_text

    from docx import Document

    docx_path = temp / "sample.docx"
    writer = Document()
    writer.add_paragraph("HERMES DOCX EXTRACTION OK")
    writer.save(docx_path)
    reader = Document(docx_path)
    checks["docx_extract"] = any(
        "HERMES DOCX EXTRACTION OK" in paragraph.text
        for paragraph in reader.paragraphs
    )

with tempfile.TemporaryDirectory(
    prefix="pytest-smoke-", dir="/opt/data/tmp"
) as raw_pytest_temp:
    pytest_temp = Path(raw_pytest_temp)
    test_file = pytest_temp / "test_runtime.py"
    test_file.write_text(
        "def test_runtime_cache():\n    assert True\n",
        encoding="utf-8",
    )
    pytest_output = run(
        "/opt/hermes/.venv/bin/python",
        "-m",
        "pytest",
        str(test_file),
        "-q",
        "--basetemp",
        str(pytest_temp / "basetemp"),
        "-o",
        "cache_dir=/opt/data/cache/pytest",
        timeout=30,
    )
    checks["pytest_targeted"] = "1 passed" in pytest_output

languages = run("tesseract", "--list-langs")
checks["ocr_arabic"] = "ara" in languages.split()
checks["hermes_admin"] = "primary=" in run("hermes-admin", "status")
quota_report = run("hermes-admin", "quota", timeout=30)
checks["quota_report"] = (
    "nararouter_free_daily_tokens=7000000" in quota_report
    and "nararouter_free_requests_per_minute=10" in quota_report
    and "nararouter_exact_remaining_tokens=dashboard-only" in quota_report
    and "local_gpu_cloud_quota=none" in quota_report
)
checks["pip_wrapper"] = "pip " in run("pip", "--version").lower()
checks["user_hermes_link"] = Path("/opt/data/.local/bin/hermes").is_symlink()
checks["gold_monitor_installed"] = Path(
    "/opt/data/home/.hermes/scripts/monitor_gold.py"
).is_file()
try:
    node_module = run(
        "node",
        "-e",
        "require('pptxgenjs'); console.log('pptxgenjs=ok')",
        timeout=10,
    )
    checks["pptxgenjs"] = "pptxgenjs=ok" in node_module
except (OSError, subprocess.SubprocessError):
    checks["pptxgenjs"] = False

with tempfile.TemporaryDirectory(
    prefix="presentation-smoke-", dir="/opt/data/tmp"
) as raw_presentation_temp:
    presentation_temp = Path(raw_presentation_temp)
    pptx_path = presentation_temp / "runtime-presentation.pptx"
    builder_path = presentation_temp / "build-presentation.cjs"
    builder_path.write_text(
        """const pptxgen = require("pptxgenjs");
const pptx = new pptxgen();
pptx.layout = "LAYOUT_WIDE";
const slide = pptx.addSlide();
slide.addText("HERMES POWERPOINT OK", {
  x: 0.8, y: 1.0, w: 11.7, h: 0.8, fontSize: 30
});
pptx.writeFile({ fileName: process.argv[2] });
""",
        encoding="utf-8",
    )
    run("node", str(builder_path), str(pptx_path), timeout=30)
    checks["presentation_create"] = (
        pptx_path.is_file() and pptx_path.stat().st_size > 0
    )
    extracted = run(
        "/opt/hermes/.venv/bin/python",
        "-m",
        "markitdown",
        str(pptx_path),
        timeout=30,
    )
    checks["presentation_extract"] = "HERMES POWERPOINT OK" in extracted
    run(
        "libreoffice",
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        str(presentation_temp),
        str(pptx_path),
        timeout=45,
    )
    rendered_pdf = presentation_temp / "runtime-presentation.pdf"
    checks["presentation_render"] = (
        rendered_pdf.is_file() and rendered_pdf.stat().st_size > 0
    )
try:
    with urllib.request.urlopen(
        "http://ollama-bridge:8000/health", timeout=10
    ) as response:
        bridge_health = json.load(response)
    checks["ollama_bridge_live"] = bool(
        bridge_health.get("status") == "ok"
        and bridge_health.get("gpu_resident") is True
    )
except (OSError, ValueError):
    checks["ollama_bridge_live"] = False

if args.network:
    from tools.web_tools import web_search_tool

    search = json.loads(web_search_tool("Open WebUI official documentation", 2))
    checks["web_search_live"] = bool(
        search.get("success") and search.get("data", {}).get("web")
    )

failed_checks = sorted(name for name, passed in checks.items() if not passed)
report = {
    "passed": not failed_checks,
    "passed_count": sum(1 for passed in checks.values() if passed),
    "check_count": len(checks),
    "failed_checks": failed_checks,
    "checks": checks,
    "primary": (
        f"{config.get('model', {}).get('provider', '')}/"
        f"{config.get('model', {}).get('default', '')}"
    ),
    "fallback_count": len(config.get("fallback_providers") or []),
}
print(json.dumps(report, indent=2, sort_keys=True))
raise SystemExit(0 if report["passed"] else 1)
